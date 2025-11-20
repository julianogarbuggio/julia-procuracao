# -*- coding: utf-8 -*-
from fastapi import FastAPI, Form, Request, HTTPException
from fastapi.responses import (
    HTMLResponse,
    FileResponse,
    PlainTextResponse,
    JSONResponse,
)
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from pathlib import Path
from datetime import datetime
from tempfile import TemporaryDirectory
import traceback
import shutil
import subprocess
import re
import unicodedata
import os
import base64

import requests
from docxtpl import DocxTemplate
from docx import Document

try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_OK = True
except Exception:
    DOCX2PDF_OK = False

# ----------------- Paths básicos -----------------
BASE_DIR = Path(__file__).resolve().parent
DOCS_DIR = BASE_DIR / "documentos"
SAIDA_DIR = BASE_DIR / "saida"
TPL_DIR = BASE_DIR / "templates"
STATIC_DIR = BASE_DIR / "static"

for p in (DOCS_DIR, SAIDA_DIR, STATIC_DIR, TPL_DIR):
    p.mkdir(parents=True, exist_ok=True)

# ----------------- Config ZapSign -----------------
ZAPSIGN_API_URL = "https://api.zapsign.com.br/api/v1/docs/"
ZAPSIGN_TOKEN = os.getenv("ZAPSIGN_API_TOKEN")  # configure no Railway

app = FastAPI(title="Jul.IA – Automação de Procuração de Revisionais de Empréstimos Consignados")
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")
templates = Jinja2Templates(directory=str(TPL_DIR))


# ----------------- Utils de parsing -----------------
def _norm(s: str) -> str:
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = re.sub(r"\s+", " ", s)
    return s.lower().strip()


def parse_bloco(texto: str) -> dict:
    """
    Lê o bloco colado (padrão WhatsApp / formulário) e devolve o contexto
    para o template DOCX.
    """
    ctx = {
        "NOME": "",
        "NACIONALIDADE": "",
        "NASCIMENTO": "",
        "ESTADO_CIVIL": "",
        "PROFISSAO": "",
        "RG": "",
        "CPF": "",
        "LOGRADOURO": "",
        "NUMERO": "",
        "COMPLEMENTO": "",
        "BAIRRO": "",
        "CEP": "",
        "CIDADE": "",
        "ESTADO": "",
        "WHATSAPP": "",
        "EMAIL": "",
        "DATA_HOJE": datetime.now().strftime("%d/%m/%Y"),
    }

    for raw in (texto or "").splitlines():
        if ":" not in raw:
            continue
        label, val = raw.split(":", 1)
        lab = _norm(label)
        value = val.strip()

        if "nome completo" in lab or lab == "nome":
            ctx["NOME"] = value
            continue
        if "nacionalidade" in lab:
            ctx["NACIONALIDADE"] = value
            continue
        if "nascimento" in lab:
            ctx["NASCIMENTO"] = value
            continue
        if "estado civil" in lab:
            ctx["ESTADO_CIVIL"] = value
            continue
        if "profiss" in lab:
            ctx["PROFISSAO"] = value
            continue

        # RG: "xxxx - ESTADO: PR" -> "xxxx - PR"
        if lab.startswith("rg"):
            num = value
            m = re.search(
                r"(?:estado\s*:\s*)?([A-Za-z]{2})\s*$",
                value,
                flags=re.IGNORECASE,
            )
            if m:
                uf = m.group(1).upper()
                num = re.sub(
                    r"-?\s*estado\s*:\s*[A-Za-z]{2}\s*$",
                    "",
                    value,
                    flags=re.IGNORECASE,
                ).strip()
                num = re.sub(r"\s*-\s*$", "", num)
                ctx["RG"] = f"{num} - {uf}"
            else:
                ctx["RG"] = value
            continue

        if "cpf" in lab:
            ctx["CPF"] = value
            continue

        # ENDEREÇO COMPLETO -> LOGRADOURO / NUMERO / COMPLEMENTO
        if "endereco" in lab or "endereço" in lab:
            mlog = re.search(r"^\s*(.*?)(?:,|$)", value)
            ctx["LOGRADOURO"] = mlog.group(1).strip() if mlog else value

            mnum = re.search(
                r"(?:^|[,;])\s*n[ºo\.]?\s*:?\s*([\d\w\-\/]+)",
                value,
                flags=re.IGNORECASE,
            )
            ctx["NUMERO"] = mnum.group(1).strip() if mnum else ""

            mcomp = re.search(
                r"complemento\s*:\s*([^,]+)",
                value,
                flags=re.IGNORECASE,
            )
            ctx["COMPLEMENTO"] = mcomp.group(1).strip() if mcomp else ""
            continue

        if "bairro" in lab:
            ctx["BAIRRO"] = value
            continue
        if "cep" in lab:
            ctx["CEP"] = value
            continue

        # CIDADE: "Maringá, ESTADO: PR"
        if "cidade" in lab and "estado" in lab:
            m = re.search(
                r"^\s*([^,\-]+)[,\-]?\s*(?:estado\s*:\s*|uf\s*:\s*)?([A-Za-z]{2})\s*$",
                value,
                flags=re.IGNORECASE,
            )
            if m:
                ctx["CIDADE"] = m.group(1).strip()
                ctx["ESTADO"] = m.group(2).upper()
            else:
                ctx["CIDADE"] = value
            continue

        if "cidade" in lab:
            m = re.search(
                r"^\s*([^,\-]+)\s*[-,]\s*(?:estado\s*:\s*|uf\s*:\s*)?([A-Za-z]{2})\s*$",
                value,
                flags=re.IGNORECASE,
            )
            if m:
                ctx["CIDADE"] = m.group(1).strip()
                ctx["ESTADO"] = m.group(2).upper()
            else:
                ctx["CIDADE"] = value
            continue

        if lab == "estado" or lab == "uf":
            ctx["ESTADO"] = value.strip().upper()
            continue

        if "whats" in lab:
            ctx["WHATSAPP"] = value
            continue
        if "e-mail" in lab or "email" in lab:
            ctx["EMAIL"] = value
            continue

    return ctx


def escolher_modelo() -> Path | None:
    preferido = DOCS_DIR / "procuracao_consignado.docx"
    if preferido.exists():
        return preferido
    for p in sorted(DOCS_DIR.glob("*.docx")):
        if "~$" in p.name:
            continue
        return p
    return None


def gerar_nome_arquivo(nome: str, extensao: str) -> Path:
    """
    Gera caminho na pasta SAIDA com o formato:
    02_Procuracao_Consig_Nome_Autor.ext
    """
    if not nome:
        sufixo = "Autor"
    else:
        sufixo = re.sub(r"\s+", "_", nome.strip())
    filename = f"02_Procuracao_Consig_{sufixo}.{extensao}"
    return SAIDA_DIR / filename


# ----------------- Negrito sem mexer em layout -----------------
def bold_nome_everywhere(docx_path: Path, nome: str):
    """
    Coloca o nome em negrito em parágrafos e tabelas,
    sem alterar estrutura ou quebras de página.
    """
    if not nome:
        return
    try:
        doc = Document(str(docx_path))

        for p in doc.paragraphs:
            for r in p.runs:
                if nome in r.text:
                    r.bold = True

        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if nome in r.text:
                                r.bold = True

        doc.save(str(docx_path))
    except Exception:
        # se der qualquer erro, não quebra o fluxo
        pass


def try_convert_with_soffice(src_docx: Path, dst_pdf: Path) -> bool:
    try:
        cmd = [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(dst_pdf.parent),
            str(src_docx),
        ]
        subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        return dst_pdf.exists()
    except Exception:
        return False


def _limpar_whatsapp(raw: str) -> str:
    """
    Remove tudo que não for dígito.
    Se vier com 55 na frente, tira o 55 e deixa só DDD+numero.
    """
    digits = re.sub(r"\D+", "", raw or "")
    if digits.startswith("55") and len(digits) > 2:
        return digits[2:]
    return digits


# ----------------- Rotas -----------------
@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/gerar-docx")
async def gerar_docx(dados: str = Form(...)):
    try:
        modelo = escolher_modelo()
        if not modelo:
            return PlainTextResponse(
                "Modelo .docx não encontrado em app/documentos.",
                status_code=500,
            )

        ctx = parse_bloco(dados)
        nome_autor = ctx.get("NOME", "").strip()

        with TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            out_docx = tmpdir / "saida.docx"

            tpl = DocxTemplate(str(modelo))
            tpl.render(ctx)
            tpl.save(str(out_docx))

            bold_nome_everywhere(out_docx, nome_autor)

            final_docx = gerar_nome_arquivo(nome_autor, "docx")
            final_docx.parent.mkdir(parents=True, exist_ok=True)
            shutil.copyfile(out_docx, final_docx)

            return FileResponse(final_docx, filename=final_docx.name)
    except Exception:
        (SAIDA_DIR / "stacktrace.txt").write_text(
            traceback.format_exc(), encoding="utf-8"
        )
        return PlainTextResponse(
            "Erro ao gerar DOCX. Veja app/saida/stacktrace.txt",
            status_code=500,
        )


@app.post("/gerar-pdf")
async def gerar_pdf(dados: str = Form(...)):
    """
    Mantido como legado / fallback.
    Usa LibreOffice no Railway e pode não ficar 100% igual ao Word.
    """
    try:
        modelo = escolher_modelo()
        if not modelo:
            return PlainTextResponse(
                "Modelo .docx não encontrado em app/documentos.",
                status_code=500,
            )

        ctx = parse_bloco(dados)
        nome_autor = ctx.get("NOME", "").strip()

        with TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            out_docx = tmpdir / "saida.docx"
            tpl = DocxTemplate(str(modelo))
            tpl.render(ctx)
            tpl.save(str(out_docx))

            bold_nome_everywhere(out_docx, nome_autor)

            out_pdf = tmpdir / "saida.pdf"
            ok = False

            if DOCX2PDF_OK:
                try:
                    docx2pdf_convert(str(out_docx), str(out_pdf))
                    ok = out_pdf.exists()
                except Exception:
                    ok = False

            if not ok:
                ok = try_convert_with_soffice(out_docx, out_pdf)

            if ok:
                final_pdf = gerar_nome_arquivo(nome_autor, "pdf")
                final_pdf.parent.mkdir(parents=True, exist_ok=True)
                shutil.copyfile(out_pdf, final_pdf)
                return FileResponse(final_pdf, filename=final_pdf.name)
            else:
                # fallback: devolve o DOCX no padrão 02_Procuracao_Consig_...
                final_docx = gerar_nome_arquivo(nome_autor, "docx")
                final_docx.parent.mkdir(parents=True, exist_ok=True)
                shutil.copyfile(out_docx, final_docx)
                return FileResponse(final_docx, filename=final_docx.name)
    except Exception:
        (SAIDA_DIR / "stacktrace.txt").write_text(
            traceback.format_exc(), encoding="utf-8"
        )
        return PlainTextResponse(
            "Erro ao tentar gerar PDF. Veja app/saida/stacktrace.txt",
            status_code=500,
        )


@app.post("/enviar-zapsign")
async def enviar_zapsign(dados: str = Form(...)):
    """
    Gera o DOCX perfeito e envia para a ZapSign criar o documento
    (PDF + fluxo de assinatura) via base64_docx.
    """
    if not ZAPSIGN_TOKEN:
        raise HTTPException(
            status_code=500,
            detail="ZAPSIGN_API_TOKEN não configurado no servidor.",
        )

    try:
        modelo = escolher_modelo()
        if not modelo:
            return PlainTextResponse(
                "Modelo .docx não encontrado em app/documentos.",
                status_code=500,
            )

        ctx = parse_bloco(dados)
        nome_autor = ctx.get("NOME", "").strip() or "Cliente"

        # 1) Gera DOCX em pasta temporária
        with TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            out_docx = tmpdir / "saida.docx"

            tpl = DocxTemplate(str(modelo))
            tpl.render(ctx)
            tpl.save(str(out_docx))

            bold_nome_everywhere(out_docx, nome_autor)

            # 2) Converte DOCX para base64 (texto)
            b64_docx = base64.b64encode(out_docx.read_bytes()).decode("ascii")

        # 3) Monta signer com dados do cliente
        signer = {
            "name": nome_autor,
        }
        email = (ctx.get("EMAIL") or "").strip()
        if email:
            signer["email"] = email

        phone = _limpar_whatsapp(ctx.get("WHATSAPP") or "")
        if phone:
            signer["phone_country"] = "55"
            signer["phone_number"] = phone

        payload = {
            "name": f"Procuração Consignado – {nome_autor}",
            "base64_docx": b64_docx,
            "lang": "pt-br",
            "signers": [signer],
            # você pode ajustar mais campos aqui (mensagens, prazo, etc.)
        }

        headers = {
            "Authorization": f"Bearer {ZAPSIGN_TOKEN}",
            "Content-Type": "application/json",
        }

        resp = requests.post(
            ZAPSIGN_API_URL, headers=headers, json=payload, timeout=40
        )

        if resp.status_code >= 300:
            return PlainTextResponse(
                f"Erro ao criar documento na ZapSign "
                f"({resp.status_code}): {resp.text}",
                status_code=500,
            )

        data = resp.json()

        # A ZapSign devolve JSON com, entre outras coisas, token e links
        # do documento (usar campo de link que vier na resposta).
        return JSONResponse(data)
    except Exception:
        (SAIDA_DIR / "stacktrace_zapsign.txt").write_text(
            traceback.format_exc(), encoding="utf-8"
        )
        return PlainTextResponse(
            "Erro ao enviar documento para ZapSign. "
            "Veja app/saida/stacktrace_zapsign.txt",
            status_code=500,
        )
